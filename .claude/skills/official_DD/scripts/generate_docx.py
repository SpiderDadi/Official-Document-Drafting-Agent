#!/usr/bin/env python3
"""
公文 Word 文档生成器
按照 GB/T 9704-2012《党政机关公文格式》标准生成规范的 Word 文档。

用法：
  python generate_docx.py --config config.json --output output.docx
  python generate_docx.py --json '{"title":"...","body":[...]}' --output output.docx
"""

import json
import argparse
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Mm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_line_spacing(paragraph, spacing_pt=28):
    """设置段落行间距为固定值（磅）"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(spacing_pt * 20)))  # 单位：1/20 磅
    spacing.set(qn('w:lineRule'), 'exact')
    pPr.append(spacing)


def set_first_line_indent(paragraph, chars=2, font_size_pt=16):
    """设置首行缩进（字符数）"""
    pPr = paragraph._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    # 首行缩进：字符数 * 字体大小（磅）* 20（twips）
    indent_twips = int(chars * font_size_pt * 20)
    ind.set(qn('w:firstLine'), str(indent_twips))
    pPr.append(ind)


def add_paragraph_with_font(doc, text, font_name='仿宋_GB2312', font_size_pt=16,
                            bold=False, alignment=None, first_line_indent=True,
                            line_spacing_pt=28):
    """添加段落并设置字体格式"""
    para = doc.add_paragraph()
    run = para.add_run(text)

    # 西文字体
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold

    # 中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

    # 对齐方式
    if alignment is not None:
        para.alignment = alignment

    # 首行缩进
    if first_line_indent:
        set_first_line_indent(para, chars=2, font_size_pt=font_size_pt)

    # 行间距
    set_line_spacing(para, spacing_pt=line_spacing_pt)

    return para


def add_blank_line(doc, font_size_pt=16, line_spacing_pt=28):
    """添加空行"""
    para = doc.add_paragraph()
    run = para.add_run('')
    run.font.size = Pt(font_size_pt)
    set_line_spacing(para, spacing_pt=line_spacing_pt)
    return para


def generate_document(config):
    """
    根据配置生成公文 Word 文档。

    config 格式：
    {
        "title": "公文标题",
        "main_recipient": "主送机关名称",
        "body": [
            {"type": "text", "content": "正文段落"},
            {"type": "heading1", "content": "一、一级标题"},
            {"type": "heading2", "content": "（一）二级标题"},
            {"type": "text", "content": "正文段落"},
            {"type": "blank", "content": ""}
        ],
        "issuing_authority": "发文机关署名",
        "date": "2026年8月24日",
        "attachments": ["附件1：XXX", "附件2：XXX"],
        "copy_to": "抄送机关名称"
    }
    """
    doc = Document()

    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)

    # ===== 标题 =====
    # 2号小标宋体（22pt），居中，加粗
    title = config.get('title', '')
    add_paragraph_with_font(
        doc, title,
        font_name='宋体', font_size_pt=22,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        line_spacing_pt=36
    )

    # 标题后空一行
    add_blank_line(doc, font_size_pt=16, line_spacing_pt=28)

    # ===== 主送机关 =====
    recipient = config.get('main_recipient', '')
    if recipient:
        # 格式：主送机关名称后加全角冒号
        add_paragraph_with_font(
            doc, recipient + '：',
            font_name='仿宋_GB2312', font_size_pt=16,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=False,
            line_spacing_pt=28
        )

    # ===== 正文 =====
    body = config.get('body', [])
    for item in body:
        item_type = item.get('type', 'text')
        content = item.get('content', '')

        if not content and item_type != 'blank':
            continue

        if item_type == 'blank':
            add_blank_line(doc, font_size_pt=16, line_spacing_pt=28)

        elif item_type == 'heading1':
            # 一级标题：3号黑体（16pt）
            add_paragraph_with_font(
                doc, content,
                font_name='黑体', font_size_pt=16,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=True,
                line_spacing_pt=28
            )

        elif item_type == 'heading2':
            # 二级标题：3号楷体（16pt），不加粗
            add_paragraph_with_font(
                doc, content,
                font_name='楷体_GB2312', font_size_pt=16,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=True,
                line_spacing_pt=28
            )

        elif item_type == 'text':
            # 正文：3号仿宋体（16pt）
            add_paragraph_with_font(
                doc, content,
                font_name='仿宋_GB2312', font_size_pt=16,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=True,
                line_spacing_pt=28
            )

    # ===== 附件说明 =====
    attachments = config.get('attachments', [])
    if attachments:
        add_blank_line(doc, font_size_pt=16, line_spacing_pt=28)
        for i, att in enumerate(attachments, 1):
            add_paragraph_with_font(
                doc, f'附件：{i}. {att}' if len(attachments) > 1 else f'附件：{att}',
                font_name='仿宋_GB2312', font_size_pt=16,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=True,
                line_spacing_pt=28
            )

    # ===== 发文机关署名和成文日期 =====
    add_blank_line(doc, font_size_pt=16, line_spacing_pt=28)

    issuing_authority = config.get('issuing_authority', '')
    if issuing_authority:
        add_paragraph_with_font(
            doc, issuing_authority,
            font_name='仿宋_GB2312', font_size_pt=16,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            first_line_indent=False,
            line_spacing_pt=28
        )

    date = config.get('date', '')
    if date:
        add_paragraph_with_font(
            doc, date,
            font_name='仿宋_GB2312', font_size_pt=16,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            first_line_indent=False,
            line_spacing_pt=28
        )

    return doc


def main():
    parser = argparse.ArgumentParser(description='公文 Word 文档生成器')
    parser.add_argument('--config', '-c', type=str, help='JSON 配置文件路径')
    parser.add_argument('--json', '-j', type=str, help='JSON 配置字符串')
    parser.add_argument('--output', '-o', type=str, required=True, help='输出 Word 文件路径')

    args = parser.parse_args()

    if args.config:
        config = json.loads(Path(args.config).read_text(encoding='utf-8'))
    elif args.json:
        config = json.loads(args.json)
    else:
        print('错误：请提供 --config 或 --json 参数', file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = generate_document(config)
    doc.save(str(output_path))
    print(f'✅ Word 文档已保存至：{output_path}')


if __name__ == '__main__':
    main()