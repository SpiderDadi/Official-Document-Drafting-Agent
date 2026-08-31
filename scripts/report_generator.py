#!/usr/bin/env python3
"""
政策深度分析报告生成器
从 AI 分析结果 JSON 生成规范的 Word 报告。

用法：
  python report_generator.py --config analysis.json --output 报告.docx
  python report_generator.py --json '{"title":"...","analysis":{...}}' --output 报告.docx

输入格式（analysis.json）：
{
  "date": "2026-08-31",
  "policies": [
    {
      "title": "政策标题",
      "source": "信源",
      "summary": "摘要",
      "analysis": {
        "overview": ["概览要点1", "概览要点2"],
        "deep_dive": ["深度要点1", "深度要点2"],
        "action_items": ["行动建议1", "行动建议2"],
        "three_rounds": [
          {"round": "第一轮：概览与框架", "qas": [{"question": "...", "answer": "..."}]},
          {"round": "第二轮：深度挖掘", "qas": [...]},
          {"round": "第三轮：综合与反刍", "qas": [...]}
        ]
      }
    }
  ]
}

设计原则：
  - 纯代码实现，0 token 消耗
  - 基于已有分析结果直接渲染
  - 遵循 GB/T 9704-2012 格式规范
"""

import json
import argparse
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("❌ 缺少 python-docx 库，请运行：pip install python-docx")
    sys.exit(1)


def set_line_spacing(paragraph, spacing_pt=28):
    """设置段落行间距（磅）"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    pPr.append(spacing)


def set_first_line_indent(paragraph, chars=2, font_size_pt=16):
    """设置首行缩进（字符数）"""
    pPr = paragraph._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    indent_twips = int(chars * font_size_pt * 20)
    ind.set(qn('w:firstLine'), str(indent_twips))
    pPr.append(ind)


def add_para(doc, text, font='仿宋_GB2312', size=16, bold=False,
             align=None, indent=True, spacing=28):
    """添加段落并设置格式"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold

    # 中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font)
    rPr.insert(0, rFonts)

    if align is not None:
        para.alignment = align

    if indent:
        set_first_line_indent(para, chars=2, font_size_pt=size)

    set_line_spacing(para, spacing_pt=spacing)
    return para


def add_blank(doc, size=16, spacing=28):
    """空行"""
    para = doc.add_paragraph()
    run = para.add_run('')
    run.font.size = Pt(size)
    set_line_spacing(para, spacing_pt=spacing)
    return para


def generate_report(config):
    """
    根据配置生成政策深度分析报告 Word 文档。

    config 格式：
    {
        "date": "2026-08-31",
        "policies": [
            {
                "title": "政策标题",
                "source": "信源",
                "summary": "摘要",
                "key_points": "关键要点",
                "rating": "★★★★★",
                "analysis": {
                    "summary": "分析总结",
                    "analysis": [
                        {"round": "第一轮...", "question": "...", "answer": "..."}
                    ],
                    "actionItems": ["...", "..."]
                }
            }
        ]
    }
    """
    doc = Document()

    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)

    date_str = config.get("date", datetime.now().strftime("%Y-%m-%d"))
    policies = config.get("policies", [])

    # ===== 封面 =====
    add_para(doc, "政策深度分析报告", font='宋体', size=22, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, spacing=36)
    add_blank(doc)
    add_para(doc, f"日期：{date_str}", font='楷体_GB2312', size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, f"重点政策 {len(policies)} 条", font='楷体_GB2312', size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_blank(doc)
    add_para(doc, "本报告由 AI 深度分析 + 自动生成系统产出",
             font='楷体_GB2312', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_blank(doc)
    add_blank(doc)

    # ===== 目录 =====
    add_para(doc, "目  录", font='黑体', size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_blank(doc)

    for i, policy in enumerate(policies, 1):
        add_para(doc, f"{i}. {policy['title']}",
                 font='仿宋_GB2312', size=16)

    add_blank(doc)
    add_blank(doc)

    # ===== 各政策分析 =====
    for i, policy in enumerate(policies, 1):
        # 政策标题
        add_para(doc, f"{'━'*30}", font='黑体', size=12, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        add_para(doc, f"{i}. {policy['title']}", font='宋体', size=18, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        add_para(doc, f"{'━'*30}", font='黑体', size=12, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        add_blank(doc)

        # 基本信息
        add_para(doc, f"信源：{policy.get('source', '未知')}",
                 font='黑体', size=16, bold=True)
        add_para(doc, f"评级：{policy.get('rating', '未评级')}",
                 font='黑体', size=16, bold=True)
        add_blank(doc)

        # 摘要
        if policy.get('summary'):
            add_para(doc, "【政策摘要】", font='黑体', size=16, bold=True)
            add_para(doc, policy['summary'])
            add_blank(doc)

        if policy.get('key_points'):
            add_para(doc, "【关键要点】", font='黑体', size=16, bold=True)
            add_para(doc, policy['key_points'])
            add_blank(doc)

        # 深度分析
        analysis = policy.get('analysis', {})
        if isinstance(analysis, dict):
            # 分析总结
            if analysis.get('summary'):
                add_para(doc, "【分析总结】", font='黑体', size=16, bold=True)
                add_para(doc, analysis['summary'])
                add_blank(doc)

            # 三轮递进分析
            rounds_data = analysis.get('analysis', [])
            if rounds_data:
                # 按轮次分组
                current_round = None
                round_items = []
                for item in rounds_data:
                    r = item.get('round', '')
                    if r != current_round:
                        if current_round and round_items:
                            _render_round(doc, current_round, round_items)
                            add_blank(doc)
                        current_round = r
                        round_items = [item]
                    else:
                        round_items.append(item)

                if current_round and round_items:
                    _render_round(doc, current_round, round_items)
                    add_blank(doc)

            # 行动建议
            action_items = analysis.get('actionItems', [])
            if action_items:
                add_para(doc, "【行动建议】", font='黑体', size=16, bold=True)
                for j, action in enumerate(action_items, 1):
                    add_para(doc, f"{j}. {action}")
                add_blank(doc)

        # 分页符（最后一个政策除外）
        if i < len(policies):
            doc.add_page_break()

    # ===== 尾注 =====
    doc.add_page_break()
    add_blank(doc)
    add_para(doc, "—— 报告结束 ——",
             font='楷体_GB2312', size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_blank(doc)
    add_para(doc, f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
             font='楷体_GB2312', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, "本报告基于 AI 深度分析结果自动生成，内容仅供参考",
             font='楷体_GB2312', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    return doc


def _render_round(doc, round_name, qa_items):
    """渲染一轮分析（Q&A 格式）"""
    add_para(doc, round_name, font='黑体', size=16, bold=True)

    for item in qa_items:
        q = item.get('question', '')
        a = item.get('answer', '')
        if q:
            add_para(doc, f"Q：{q}", font='楷体_GB2312', size=16, bold=True)
        if a:
            add_para(doc, f"A：{a}")


def main():
    parser = argparse.ArgumentParser(description="政策深度分析报告生成器")
    parser.add_argument('--config', '-c', type=str,
                        help='JSON 配置文件路径')
    parser.add_argument('--json', '-j', type=str,
                        help='JSON 配置字符串')
    parser.add_argument('--output', '-o', type=str, required=True,
                        help='输出 Word 文件路径')

    args = parser.parse_args()

    if args.config:
        config = json.loads(Path(args.config).read_text(encoding="utf-8"))
    elif args.json:
        config = json.loads(args.json)
    else:
        print("错误：请提供 --config 或 --json 参数", file=sys.stderr)
        sys.exit(1)

    # 如果没有 date，自动填充
    if "date" not in config:
        config["date"] = datetime.now().strftime("%Y-%m-%d")

    # 如果没有 policies，尝试从旧格式转换
    if "policies" not in config:
        # 兼容旧格式：直接是单个政策的分析
        if "title" in config:
            config["policies"] = [config]
        elif "analyses" in config:
            config["policies"] = config["analyses"]

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = generate_report(config)
    doc.save(str(output_path))
    print(f"✅ 政策深度分析报告已保存：{output_path}")


if __name__ == "__main__":
    main()
